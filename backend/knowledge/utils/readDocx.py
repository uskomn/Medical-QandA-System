from docx import Document

def readDocx(file_path):
    doc=Document(file_path)
    full_text=[]
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    document='\n'.join(full_text)
    return document


